from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "功能一任务一_场景建模与数据接口设计_提交说明报告.docx"
BLUE, DARK, GRAY = "2E74B5", "1F4D78", "666666"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def set_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None: tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]; set_width(cell, widths[i]); shade(cell, "E8EEF5")
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); r = p.add_run(label); set_font(r, 10, True, DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_width(cells[i], widths[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2); r = p.add_run(str(value)); set_font(r, 9.5)
    for row in table.rows:
        for cell in row.cells:
            cell.margin_top = 80; cell.margin_bottom = 80; cell.margin_left = 120; cell.margin_right = 120
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16 if level == 1 else 10); p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    r = p.add_run(text); set_font(r, 16 if level == 1 else 13, True, BLUE if level == 1 else DARK)


def add_body(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), 11)


doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)
styles = doc.styles
styles["Normal"].font.name = "Calibri"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); styles["Normal"].font.size = Pt(11)

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_font(header.add_run("功能一·任务一｜提交说明报告"), 9, False, GRAY)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(footer.add_run("面向雄安新区“城市大脑”的车路云一体化协同管控算法与仿真平台研究"), 8.5, False, GRAY)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(36); p.paragraph_format.space_after = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run("功能一·任务一"), 18, True, BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(12); set_font(p.add_run("场景建模与数据接口设计\n提交说明报告"), 26, True, DARK)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(28); set_font(p.add_run("OSM-SUMO 场景生成与仿真运行时数据接口"), 13, False, GRAY)
add_table(doc, ["项目", "说明"], [["对应任务", "功能一·任务一：场景建模与数据接口设计"], ["交付目录", "osm-api/"], ["场景范围", "雄安新区坐标示例（38.9429, 115.8977），可由 API 动态生成"], ["报告日期", "2026-08-28"]], [1800, 7560])
add_body(doc, "本报告说明任务一交付包的模块组成、接口边界、复现方法及验收状态。报告中的“示例”仅用于字段和接口联调，不替代真实 SUMO 运行实验结果。")

add_heading(doc, "1. 任务目标与交付范围")
add_body(doc, "任务一面向交通协同控制算法提供可复用的场景建模与数据接口基础：根据目标区域坐标生成 OSM-SUMO 场景，并以统一 HTTP 接口支持路口状态读取、信号动作下发、仿真步进和会话关闭。算法实现、性能对比和真实场景结论不属于本任务的接口层交付。")
add_table(doc, ["交付项", "实现位置", "用途"], [["场景生成", "src/service.py", "OSMnx 下载、netconvert 转网、randomTrips 生成车流并打包"], ["运行时接口", "src/service.py", "启动、状态、动作、步进、关闭"], ["接口契约", "contracts/", "固定请求、观测和动作字段及合法范围"], ["文档与验证", "docs/、scripts/verify_delivery.py", "支持联调、复现和静态验收"]], [1800, 3000, 4560])

add_heading(doc, "2. 场景与接口设计")
add_body(doc, "场景生成请求接受经纬度、范围和仿真时长。成功后返回自包含 ZIP，内含 network.net.xml、traffic.rou.xml 和 scenario.sumocfg。运行时会话从 .sumocfg 启动，并以 SUMO 生成的 traffic-light ID 作为路口唯一标识，避免在算法端写死节点编号。")
add_table(doc, ["接口", "关键输入", "关键输出"], [["POST /v1/scenario/generate", "latitude、longitude、radius_m、duration_s", "SUMO 场景 ZIP"], ["POST /v1/simulation/start", "config_path", "intersection_ids"], ["GET /v1/simulation/state", "可选 intersection_id", "相位、队列、车道车辆数、速度、占有率"], ["POST /v1/simulation/action", "intersection_id、phase、duration_s", "已接受动作"], ["POST /v1/simulation/step", "steps", "推进后的 time_s"]], [2500, 3300, 3560])
add_body(doc, "完整字段字典、错误语义、PowerShell 调用示例和“状态→动作→步进”时序已写入 docs/API.md 与 docs/INTERFACE_PROTOCOL.md。服务未启动会话时返回 HTTP 409；请求字段不符合契约时返回 400 或 404。")

add_heading(doc, "3. 工程化与可复现性")
add_body(doc, "交付包采用与功能二任务二相同的分层组织：src 放置实现，configs 保存默认场景参数，contracts 固定接口约束，data 放置可审计示例，docs 说明架构与验收，scripts 提供独立检查入口。原始 scenario/ 文件予以保留；其临时名称也保留，以防破坏 .sumocfg 引用。")
add_table(doc, ["验收动作", "命令/依据", "预期结果"], [["静态完整性", "python scripts/verify_delivery.py", "目录、文档和 JSON 契约均可读取"], ["服务健康", "GET /health", "返回服务状态"], ["场景元数据", "GET /v1/scenario/metadata", "返回默认坐标、范围、时长"], ["SUMO 闭环", "start → state → action → step → close", "返回实际路口 ID 与连续时间步状态"]], [2100, 3800, 3460])

add_heading(doc, "4. 本机验证边界与后续验收")
add_body(doc, "2026-08-28 对原始 sim_api_service.py 的直接启动检查显示，本机 Python 环境缺少 Flask，服务未能进入启动阶段。因此本次完成的是代码、接口契约和不依赖外部环境的交付完整性改造；未声明或虚构 SUMO 动态运行、OSM 下载或算法性能结果。")
add_body(doc, "部署人员应在 Python 3.10+ 虚拟环境中执行 pip install -r requirements.txt，配置 SUMO_HOME，并确认本机可访问 OpenStreetMap。随后使用真实 SUMO 路口 ID 完成一次闭环调用，将终端输出、路口状态 JSON 与运行截图存入 outputs/，作为报告和接口字典的运行证据。")
add_heading(doc, "5. 结论")
add_body(doc, "交付包已将原有的单文件接口雏形扩展为符合功能一任务一要求的场景建模与数据接口模块：具备 OSM-SUMO 场景生成入口、仿真状态与动作接口、JSON Schema、字段字典、时序说明、样例数据、复现说明和验收清单。动态仿真验证依赖本机 SUMO 与 Python 依赖环境，需在环境补齐后按本文第 3 节闭环完成。")

doc.core_properties.title = "功能一任务一 场景建模与数据接口设计 提交说明报告"
doc.core_properties.author = "User"
doc.save(OUT)
print(OUT)
