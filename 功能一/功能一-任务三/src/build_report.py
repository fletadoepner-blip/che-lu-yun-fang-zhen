"""Create the Function-2 platform report from generated, auditable artifacts."""
from __future__ import annotations
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]; DATA=ROOT/'data'/'xiong_an_20'; OUT=ROOT/'reports'/'雄安新区20路口高保真仿真验证平台报告.docx'
def cell(cell, text, bold=False):
    cell.text=str(text); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=bold
def heading(doc,text,level=1):
    p=doc.add_paragraph();p.style=f'Heading {level}';p.add_run(text)
def main():
    manifest=json.loads((DATA/'manifest.json').read_text(encoding='utf-8')); roadnet=json.loads((DATA/'roadnet.json').read_text(encoding='utf-8'))
    totals={p:sum(x['periods'][p]['vehicles'] for x in manifest) for p in ('morning','midday','evening')}
    doc=Document(); sec=doc.sections[0];sec.top_margin=Cm(2.2);sec.bottom_margin=Cm(2);sec.left_margin=Cm(2.2);sec.right_margin=Cm(2.2)
    styles=doc.styles; styles['Normal'].font.name='Microsoft YaHei';styles['Normal'].font.size=Pt(10.5)
    title=doc.add_paragraph();title.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=title.add_run('雄安新区容东片区 20 路口高保真仿真验证平台报告');r.bold=True;r.font.name='Microsoft YaHei';r.font.size=Pt(20);r.font.color.rgb=RGBColor(25,76,122)
    sub=doc.add_paragraph('赛题 XH-202613 功能二：高保真仿真验证平台的通用性架构搭建');sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('版本：1.0    生成日期：2026-08-16    数据来源：题目提供的 20 份路口流量和配时工作簿。')
    heading(doc,'1. 建设目标')
    doc.add_paragraph('平台面向雄安新区容东片区“窄路密网”控制验证，提供可复现的 20 路口联动 CityFlow 场景、统一算法接入、运行指标采集和动态可视化。路口标签来自随附高德截图文件名、地标和已有路口 2 位置说明；能确认的道路/门牌写入名称，不能确认的节点明确标为近似位置。模型拓扑依据这些材料整理为带轻微弧度的近似 4×5 网络，不能替代 OSM 或 netedit 的测绘级车道模型。')
    heading(doc,'2. 场景与数据')
    virtual_count=sum(1 for item in roadnet['intersections'] if item['virtual'])
    doc.add_paragraph(f'路网包含 20 个非虚拟信号控制节点、{len(roadnet["roads"])} 条道路和 {virtual_count} 个边界接口。每个路口保留 Excel 记录的东、西、南、北及必要的东北/西北/东南/西南进口左转、直行、右转需求；相邻方向可通行时，车辆继续穿过相邻节点，从而形成可控的交通传播。')
    t=doc.add_table(rows=1, cols=4);t.style='Table Grid'
    for c,v in zip(t.rows[0].cells,['时段','车辆数','生成窗口','车辆组成']):cell(c,v,True)
    for period,label in [('morning','早高峰 07:00-09:00'),('midday','平峰 14:30-16:30'),('evening','晚高峰 17:30-19:30')]:
        row=t.add_row().cells;cell(row[0],label);cell(row[1],f'{totals[period]:,}');cell(row[2],'0-7,199 s');cell(row[3],'小客车 86%，公交 8%，货运 6%（可复现抽样）')
    doc.add_paragraph('原始配时表逐项存储在 data/xiong_an_20/source_signal_plans.json。由于原表相位数、相位文字和周期不一致，联动算法使用统一八相位保护左转/直右动作集；原始方案作为固定配时基线和审计材料保留，未被伪造为统一训练标签。')
    heading(doc,'3. 平台架构')
    doc.add_paragraph('数据层读取 Excel 并生成 CityFlow roadnet/flow；仿真层负责多线程 CityFlow 运行；适配层以 SignalAlgorithm 契约加载固定配时、最大压力和 UGAT+FRAP；采集层写出 CSV/JSON 指标；展示层将轨迹转为无需服务端的 HTML 仪表盘。')
    heading(doc,'4. 算法适配与优化')
    doc.add_paragraph('UGAT 基座参数由 frozen_ugat_4x4.pt 加载后强制 requires_grad=False，FRAP 关系网络和融合标量是唯一可训练参数。20 个路口状态被合为批张量，推理使用 torch.inference_mode()，减少逐路口框架开销。流量生成使用单次解析、固定随机种子和排序写入，保证总量审计与复验。')
    heading(doc,'5. 可复现部署')
    doc.add_paragraph('Dockerfile 在镜像构建阶段从项目内固定的官方 CityFlow 源码（third_party/CityFlow，提交 81ee0f47659ca66177a71f81676691c58ee89184）编译 Python 扩展，并运行静态场景校验；docker-compose 默认启动最大压力基线。该源码构建镜像已实际执行 CityFlow 回归，不依赖基础镜像预装的 CityFlow 二进制。')
    heading(doc,'6. 验收结果与运行步骤')
    doc.add_paragraph('已通过静态验收：20 个受控路口、所有信号节点 8 个动作加 1 个清空相位、每一个相邻道路转接均存在 CityFlow roadLink、三时段流量总数与 20 份工作簿解析汇总一致。excel_audit.json 已逐表核验 20/20 份工作簿的进口方向和车辆总量；斜向进口按 Excel 原始标签保留，并在可视化中折入主路走廊。容器内最终路网的早高峰 7,500 步完整回归已实际运行：77,749 辆需求、40.319 秒、186.02 step/s、结束在网 84 辆、排队代理量 142。该记录验证平台可运行；新路网完成后尚未重新完成最大压力多种子基线对比，因此不能将 FRAP 宣称为性能提升。运行顺序：1) python src/build_xiong_an_20.py；2) python src/validate_scenario.py；3) docker compose up --build；4) powershell -ExecutionPolicy Bypass -File scripts/run_and_show.ps1。')
    heading(doc,'7. 边界与后续工作')
    doc.add_paragraph('若需要满足“基于 OSM/netedit”的严格车道级验证，应从高德/OSM 合规数据源补充经纬度、道路等级、车道数、限速、匝道和真实节点连接，并用 netedit 复核。本交付已提供可替换的 CityFlow 数据层和算法契约，可在不改动算法接入层的前提下替换路网。')
    heading(doc,'8. 复现操作指南')
    doc.add_paragraph('运行前准备：Windows 10/11、Docker Desktop 已启动，命令行当前目录为项目根目录。不要在 third_party/CityFlow 中单独运行脚本；该目录由 Dockerfile 在构建镜像时编译。若仅复现随项目交付的场景，不需要重新输入 Excel 数据。只有修改源数据时，才需要输入题目工作簿目录。')
    doc.add_paragraph('步骤 1，检查场景文件：在 PowerShell 输入：python .\\src\\validate_scenario.py。预期输出必须包含 status=PASS、controlled_intersections=20，并显示 morning=77749、midday=54864、evening=87051。任何 AssertionError 或非零退出码均表示数据、路由或文件不完整，不能进入算法对比。')
    doc.add_paragraph('步骤 2，构建官方 CityFlow 镜像：输入：docker build -t xiong-an-20-platform:final . 。预期日志包含 Successfully built CityFlow 和 Successfully installed CityFlow；这表示当前项目内 fixed CityFlow 源码已完成 C++/Python 扩展编译。')
    doc.add_paragraph('步骤 3，验证引擎与 100 步冒烟测试：输入：docker run --rm -v "${PWD}:/workspace/final" --entrypoint /bin/bash xiong-an-20-platform:final -lc "cd /workspace/final && python src/check_cityflow.py && python src/run_cityflow.py --period morning --algorithm ugat_frap --steps 100 --threads 1"。预期 check_cityflow 输出 engine=true，随后输出一条 algorithm=ugat_frap 的 JSON。日志中不能出现 Invalid route、load config failed 或 Traceback。')
    doc.add_paragraph('步骤 4，正式对比。请在 PowerShell 中先执行最大压力基线。以下五行必须一起复制执行：')
    doc.add_paragraph('docker run --rm `\n  -v "${PWD}:/workspace/final" `\n  --entrypoint /bin/bash `\n  xiong-an-20-platform:final `\n  -lc "cd /workspace/final && python src/run_cityflow.py --period morning --algorithm max_pressure --steps 7500 --threads 4"')
    doc.add_paragraph('基线完成后，执行 UGAT+FRAP。命令完全相同，只把 --algorithm max_pressure 改为 --algorithm ugat_frap：')
    doc.add_paragraph('docker run --rm `\n  -v "${PWD}:/workspace/final" `\n  --entrypoint /bin/bash `\n  xiong-an-20-platform:final `\n  -lc "cd /workspace/final && python src/run_cityflow.py --period morning --algorithm ugat_frap --steps 7500 --threads 4"')
    doc.add_paragraph('两次必须使用相同的 period、steps 和 threads，结果才可比较。输出会追加写入 outputs/metrics.csv；两条轨迹分别写入 outputs/trace_max_pressure_morning.json 与 outputs/trace_ugat_frap_morning.json。')
    doc.add_paragraph('步骤 5，生成可视化：输入：python .\\src\\dashboard.py .\\outputs\\trace_ugat_frap_morning.json --out .\\outputs\\dashboard.html，然后用浏览器打开 outputs/dashboard.html。若要重新由 Excel 生成三时段流量，输入：python .\\src\\build_xiong_an_20.py --source-dir "C:\\Users\\lauri\\Desktop\\路口数据\\路口数据"，随后必须重新执行步骤 1 至步骤 4。')
    doc.add_paragraph('不使用浏览器的一键动态展示方式：输入：powershell -ExecutionPolicy Bypass -File .\\scripts\\run_and_show.ps1 -Algorithm ugat_frap -Period morning -Steps 7500 -Threads 4。该脚本会自动检查或构建 Docker 镜像，在 CityFlow 运行开始后立即弹出 Windows 原生动态窗口。窗口实时显示 4×5 路网、20 个信号相位、节点排队量和 CityFlow 返回的车辆道路位置；仿真结束后保留最终状态。比较基线时把 -Algorithm ugat_frap 改为 -Algorithm max_pressure。')
    heading(doc,'9. 动态窗口数据与颜色说明')
    doc.add_paragraph('路口名称显示在节点周边的空白区域，避免覆盖车道。窗口底部字段均为 CityFlow 当前帧数据，不是累计性能结论。')
    display=doc.add_table(rows=1, cols=2);display.style='Table Grid'
    for c,v in zip(display.rows[0].cells,['窗口字段','含义']):cell(c,v,True)
    for values in [
        ('simulation time','当前仿真时刻，单位秒。'),
        ('on-network','当前时刻仍在 CityFlow 路网内的车辆数，不是总需求。'),
        ('injected=a/b','a 为已按 Excel 时段调度进入仿真的车辆数；b 为该时段总需求车辆数。'),
        ('queued','20 个路口入口车道车辆数的合计代理指标；各路口标签的 queue 为该节点对应值。'),
        ('displayed','为保障窗口刷新而显示的车辆点数量；点位来自 CityFlow 车辆道路与行驶距离。'),
        ('RUNNING / COMPLETE','分别表示仿真或轨迹回放仍在进行、已完成。')]:
        row=display.add_row().cells
        for c,v in zip(row,values):cell(c,v)
    doc.add_paragraph('路口圆点颜色表示 UGAT+FRAP 在该帧为该路口选择的信号放行相位，不表示拥堵等级、优劣或风险颜色。')
    phases=doc.add_table(rows=1, cols=3);phases.style='Table Grid'
    for c,v in zip(phases.rows[0].cells,['窗口颜色','动作编号','当前放行相位']):cell(c,v,True)
    for values in [
        ('深绿','0','东西直行及右转'), ('亮绿','1','南北直行及右转'), ('黄绿','2','东西保护左转'), ('黄褐','3','南北保护左转'),
        ('橙','4','西向左转及直行/右转'), ('红','5','东向左转及直行/右转'), ('紫','6','南向左转及直行/右转'), ('青蓝','7','北向左转及直行/右转')]:
        row=phases.add_row().cells
        for c,v in zip(row,values):cell(c,v)
    heading(doc,'10. 结果判定标准')
    criteria=doc.add_table(rows=1, cols=3);criteria.style='Table Grid'
    for c,v in zip(criteria.rows[0].cells,['检查维度','合格标准','优秀/可宣称优化标准']):cell(c,v,True)
    for values in [
        ('数据与路由','静态校验 PASS；20 个控制路口；无 Invalid route。','三时段均 PASS，重建后车辆总数与清单完全一致。'),
        ('CityFlow 可运行性','check_cityflow 返回 engine=true；100 步退出码为 0。','正式 7,500 步运行无 Traceback、无配置/路由警告。'),
        ('算法公平比较','相同流量时段、步数、线程数、随机种子。','至少连续 3 个随机种子/时段重复，报告均值和标准差。'),
        ('拥堵效果','仅记录 final_queue_proxy 和 final_active_vehicles，不作优越性结论。','UGAT+FRAP 相对最大压力基线，两个指标均不增加，且排队代理量平均降低至少 5%。'),
        ('运行效率','记录 steps_per_second；机器配置同时记录。','在相同硬件和线程数下，UGAT+FRAP 速度不低于基线的 80%。')]:
        row=criteria.add_row().cells
        for c,v in zip(row,values):cell(c,v)
    doc.add_paragraph('当前已记录的最终路网早高峰 7,500 步 UGAT+FRAP 结果为：77,749 辆输入需求、结束在网 84 辆、结束排队代理量 142、186.02 step/s。该结果达到“场景、引擎和算法可复现运行”的合格标准。由于斜向进口修正后尚未重新完成最大压力的多种子基线对比，当前不能对 FRAP 作性能优越性结论；需以本节相同条件完成基线、重复种子和适配器训练后再更新结论。')
    OUT.parent.mkdir(parents=True,exist_ok=True);doc.save(OUT);print(OUT)
if __name__=='__main__':main()
