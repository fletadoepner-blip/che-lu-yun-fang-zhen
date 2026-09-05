from __future__ import annotations
import csv
import json
from pathlib import Path
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "平峰压力扰动试验"
LEVELS = [(10, "+10%"), (20, "+20%"), (30, "+30%")]
METHODS = [("固定配时", "固定配时"), ("UGAT_FRAP_TRANSYT", "UGAT+FRAP+TRANSYT")]
CANONICAL = ["algorithm", "period", "steps", "wall_seconds", "steps_per_second", "total_demand", "scheduled_vehicles", "completed_vehicles_est", "throughput_est", "final_active_vehicles", "final_queue_proxy", "average_travel_time_s", "estimated_delay_s", "freeflow_mean_s"]

def read(path: Path) -> dict[str, float | str]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        rows = list(csv.reader(f))
    vals = dict(zip(rows[0], rows[1])) if rows and rows[0] and rows[0][0] == "algorithm" else dict(zip(CANONICAL, rows[-1]))
    out: dict[str, float | str] = {}
    for k, v in vals.items():
        try: out[k] = float(v)
        except ValueError: out[k] = v
    trace = next(path.parent.glob("trace_*.json"))
    records = json.loads(trace.read_text(encoding="utf-8"))
    queues = [float(item.get("queue_proxy", 0)) for item in records]
    out["mean_queue_proxy"] = sum(queues) / len(queues) if queues else 0.0
    return out

def cell(c, value, bold=False, color=None):
    c.text = str(value); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Microsoft YaHei"; r.font.size = Pt(9); r.bold = bold
            if color: r.font.color.rgb = RGBColor(*color)

def main():
    rows = []
    for level, label in LEVELS:
        for method, name in METHODS:
            r = read(OUT / f"压力{level}" / method / "metrics.csv")
            r.update(level=level, level_label=label, method=method, method_label=name)
            rows.append(r)
    fields = ["increase_pct", "algorithm", "total_demand", "average_travel_time_s", "estimated_delay_s", "mean_queue_proxy", "throughput_est", "completed_vehicles_est"]
    with (OUT / "压力梯度_固定配时_vs_UGAT_FRAP_TRANSYT汇总.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields); w.writeheader()
        for r in rows: w.writerow({"increase_pct": r["level"], "algorithm": r["method_label"], **{k: r.get(k, "") for k in fields[2:]}})

    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]; plt.rcParams["axes.unicode_minus"] = False
    charts = []
    for metric, ylabel, title, file in [("average_travel_time_s", "平均旅行时间 (s)", "平均旅行时间对比", "压力梯度_折线图_平均旅行时间.png"), ("estimated_delay_s", "估计延误 (s)", "估计延误对比", "压力梯度_折线图_估计延误.png"), ("mean_queue_proxy", "全过程平均队列代理", "全过程平均队列代理对比", "压力梯度_折线图_平均队列代理.png")]:
        fig, ax = plt.subplots(figsize=(9.2, 5.4)); x = list(range(3));
        for idx, (method, name) in enumerate(METHODS):
            vals = [float(next(r[metric] for r in rows if r["method"] == method and r["level"] == level)) for level, _ in LEVELS]
            color = ["#6b7280", "#1677b8"][idx]; line, = ax.plot(x, vals, marker="o", linewidth=2.5, markersize=8, color=color, label=name)
            for i, v in enumerate(vals):
                offset = 12 if idx == 0 else -18
                x_shift = -0.045 if idx == 0 else 0.045
                ax.annotate(f"{v:.1f}", (x[i] + x_shift, v), xytext=(0, offset), textcoords="offset points", ha="center", va="bottom" if offset > 0 else "top", fontsize=9, color=color, bbox=dict(boxstyle="round,pad=0.18", fc="white", ec=color, alpha=.9))
        ymin, ymax = ax.get_ylim(); ax.set_ylim(ymin - (ymax - ymin) * 0.08, ymax + (ymax - ymin) * 0.08)
        ax.set_xticks(x, [label for _, label in LEVELS]); ax.set_xlabel("平峰需求压力"); ax.set_ylabel(ylabel); ax.set_title(f"雄安新区20路口：{title}", weight="bold"); ax.grid(axis="y", alpha=.25); ax.legend(frameon=False, loc="best"); fig.tight_layout(); path = OUT / file; fig.savefig(path, dpi=220, bbox_inches="tight"); plt.close(fig); charts.append(path)

    doc = Document(); sec = doc.sections[0]; sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.75); sec.right_margin = Inches(.75)
    normal = doc.styles["Normal"]; normal.font.name = "Microsoft YaHei"; normal.font.size = Pt(10)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = title.add_run("平峰需求压力扰动实验报告"); run.bold = True; run.font.size = Pt(22); run.font.color.rgb = RGBColor(22, 77, 120)
    sub = doc.add_paragraph("固定配时 vs UGAT+FRAP+TRANSYT | 雄安新区20路口 CityFlow"); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("1. 实验目的与设置", 1)
    doc.add_paragraph("在相同雄安新区20路口路网和控制参数下，分别将平峰交通需求提高10%、20%、30%，比较固定配时基线与完整 UGAT+FRAP+TRANSYT 协同控制器的运行效果。每组运行7500步，决策间隔10步，单线程；除控制算法和压力流量文件外，其余条件保持一致。")
    doc.add_heading("2. 结果总览", 1)
    t = doc.add_table(rows=1, cols=7); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["压力", "算法", "需求量", "平均旅行时间(s)", "估计延误(s)", "全过程平均队列代理", "吞吐率"]
    for c, h in zip(t.rows[0].cells, headers): cell(c, h, True)
    for level, label in LEVELS:
        for method, name in METHODS:
            r = next(row for row in rows if row["level"] == level and row["method"] == method); vals = [label, name, f"{r['total_demand']:.0f}", f"{r['average_travel_time_s']:.3f}", f"{r['estimated_delay_s']:.3f}", f"{r['mean_queue_proxy']:.3f}", f"{float(r['throughput_est'])*100:.4f}%"]
            cs = t.add_row().cells
            for c, v in zip(cs, vals): cell(c, v)
            if method == "UGAT_FRAP_TRANSYT":
                for c in cs: c._tc.get_or_add_tcPr().append(__import__('docx').oxml.OxmlElement('w:shd')); c._tc.tcPr[-1].set(__import__('docx').oxml.ns.qn('w:fill'), 'E8F3FA')
    doc.add_heading("3. 三种压力下的折线图", 1)
    for path in charts:
        p = doc.add_paragraph(path.stem.replace("压力梯度_折线图_", "")); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_picture(str(path), width=Inches(6.7)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("4. 结论", 1)
    doc.add_paragraph("数据核查说明：此前交错插入增量车辆的压力流量文件会改变 CityFlow 车辆随机车道选择的调用顺序，导致相同 seed 下不同压力等级的基线车辆不再完全同分布，并出现旅行时间随压力下降的非物理现象。以下结论使用修正后的有序压力流量文件，旧结果不再引用。")
    for level, label in LEVELS:
        fixed = next(r for r in rows if r["level"] == level and r["method"] == "固定配时"); full = next(r for r in rows if r["level"] == level and r["method"] == "UGAT_FRAP_TRANSYT")
        gain = (float(fixed["average_travel_time_s"]) - float(full["average_travel_time_s"])) / float(fixed["average_travel_time_s"]) * 100
        doc.add_paragraph(f"{label}：完整协同平均旅行时间 {float(full['average_travel_time_s']):.3f} s，固定配时 {float(fixed['average_travel_time_s']):.3f} s，相对固定配时改善 {gain:.2f}%。", style="List Bullet")
    doc.add_paragraph("图中每个数据点均采用白底边框标签，并按算法上下错位标注；标签不覆盖数据点或另一条曲线。全过程平均队列代理由每组原始轨迹 JSON 的全部采样点求平均，避免末时刻整数快照造成两条曲线看起来相同。")
    doc.add_heading("5. 原始证据与复现", 1)
    doc.add_paragraph("原始结果位于 平峰压力扰动试验/压力10、压力20、压力30 下的固定配时和 UGAT_FRAP_TRANSYT 子目录，包含 metrics.csv、trace JSON、live JSONL 和运行配置。压力流量文件保留基线记录顺序、仅在末尾追加增量车辆，避免新增车辆改变 CityFlow 的随机车道选择序列。重跑命令：powershell -ExecutionPolicy Bypass -File ..\\scripts\\run_pressure_transyt.ps1；复算并生成本报告：python ..\\scripts\\build_pressure_report_transyt.py。")
    doc.save(OUT / "平峰压力扰动_固定配时_vs_UGAT_FRAP_TRANSYT实验报告.docx")

if __name__ == "__main__": main()
