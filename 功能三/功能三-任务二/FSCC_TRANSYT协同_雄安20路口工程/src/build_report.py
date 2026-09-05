from pathlib import Path
import csv, json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]; DATA=ROOT/'data'/'xiong_an_20'; OUT=ROOT/'reports'/'雄安新区20路口_TRANSYT高保真仿真报告.docx'
def cell(c,v,bold=False):
    c.text=str(v); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=bold
def main():
    manifest=json.loads((DATA/'manifest.json').read_text(encoding='utf-8')); roadnet=json.loads((DATA/'roadnet.json').read_text(encoding='utf-8'))
    rows=[]; path=ROOT/'outputs'/'metrics.csv'
    if path.exists():
        with path.open(encoding='utf-8',newline='') as f: rows=list(csv.DictReader(f))
    transyt=[r for r in rows if r.get('algorithm') in ('transyt','transyt_style','ugat_frap_transyt')]
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(2); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.2)
    doc.styles['Normal'].font.name='Microsoft YaHei'; doc.styles['Normal'].font.size=Pt(10.5)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('雄安新区 20 路口 TRANSYT 风格控制仿真报告'); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=RGBColor(25,76,122)
    p=doc.add_paragraph('UGAT 冻结框架 + CityFlow | 算法替换版 | 生成日期：2026-08-24'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1. 项目范围',1); doc.add_paragraph('本项目复制原 UGAT+FRAP 工程，仅替换信号控制算法，保留 roadnet、三时段 flow、拓扑、原始信号方案和路网图像。仿真入口为 src/run_cityflow.py，默认决策间隔 10 步，正式验收运行 7,500 步。')
    doc.add_heading('2. TRANSYT 风格实现与深度协同',1); doc.add_paragraph('TRANSYT 商业软件并非可直接嵌入的开源代码。本项目采用可审计的开源研究复现：以进口车道需求估计相位流量，使用 Webster 风格需求加权绿信比分配公共 90 秒周期，并依据 5×4 路网自由流传播距离设置走廊偏移；需求使用指数平滑并保留当前需求响应，每 30 个决策步重新估计一次方案。UGAT+FRAP+TRANSYT 协同模式中，TRANSYT 保持网络级周期与偏移先验；冻结 UGAT+FRAP 只在局部压力增益至少 5%、TRANSYT 当前相位低效并满足最小绿灯约束时接管。每个路口每个决定周期只下发一个动作。')
    doc.add_heading('3. 场景与数据审计',1); controlled=sum(1 for x in roadnet['intersections'] if not x.get('virtual',False)); totals={p:sum(x['periods'][p]['vehicles'] for x in manifest) for p in ('morning','midday','evening')}; doc.add_paragraph(f'路网包含 {controlled} 个非虚拟信号路口、{len(roadnet["roads"])} 条道路。原始数据文件按字节复制到本项目 data/xiong_an_20/，未重新采样或修改。')
    t=doc.add_table(rows=1,cols=3); t.style='Table Grid'
    for c,v in zip(t.rows[0].cells,('时段','需求车辆数','输入文件')): cell(c,v,True)
    for period,label in (('morning','早高峰'),('midday','平峰'),('evening','晚高峰')):
        cs=t.add_row().cells; cell(cs[0],label); cell(cs[1],f'{totals[period]:,}'); cell(cs[2],f'flow_{period}.json')
    doc.add_heading('4. 7,500 步运行结果',1)
    if transyt:
        t=doc.add_table(rows=1,cols=11); t.style='Table Grid'; fields=[('period','时段'),('steps','步数'),('total_demand','总需求'),('completed_vehicles_est','完成车辆估计'),('throughput_est','吞吐率'),('average_travel_time_s','平均旅行时间(s)'),('estimated_delay_s','估计延误(s)'),('final_queue_proxy','最终排队代理量'),('final_active_vehicles','结束在网车辆'),('frap_override_rate','FRAP接管率'),('transyt_frap_agreement_rate','动作一致率')]
        for c,(_,v) in zip(t.rows[0].cells,fields): cell(c,v,True)
        for row in transyt:
            cs=t.add_row().cells
            for c,(k,_) in zip(cs,fields): cell(c, f'{float(row.get(k,0)):.4f}' if k=='throughput_est' else row.get(k,''))
        doc.add_paragraph('average_travel_time_s 为 CityFlow 平均旅行时间；estimated_delay_s 为其与自由流时间均值之差；final_queue_proxy 为 20 个路口入口车道车辆数合计；throughput_est 为完成车辆估计比例。')
    else: doc.add_paragraph('尚未检测到协同记录。请执行 python src/run_cityflow.py --period morning --algorithm ugat_frap_transyt --steps 7500 --threads 1。')
    doc.add_heading('5. 复现命令',1); doc.add_paragraph('python src/validate_scenario.py\npython src/run_cityflow.py --period morning --algorithm ugat_frap_transyt --steps 7500 --threads 1\npython src/build_report.py'); doc.add_paragraph('三时段运行可将 --period 改为 midday 或 evening；结果追加写入 outputs/metrics.csv，轨迹写入 outputs/trace_ugat_frap_transyt_<period>.json。')
    doc.add_heading('6. 限制与解释',1); doc.add_paragraph('queue 为入口车道计数代理量，不等同于逐车排队长度；完成车辆为基于调度数和在网数的估计值。TRANSYT 风格控制器是确定性离线配时基线，应在相同流量、步数和线程条件下与 FRAP 比较，不能据单次运行宣称普遍优于 FRAP。')
    OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
if __name__=='__main__': main()
